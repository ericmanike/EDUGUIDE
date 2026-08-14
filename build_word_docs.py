import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_borders(cell, top="CCCCCC", bottom="CCCCCC", left="CCCCCC", right="CCCCCC", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{top}"/>\n'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{left}"/>\n'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{bottom}"/>\n'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{right}"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def format_paragraph(p, space_before=0, space_after=6, line_spacing=1.15):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    format_paragraph(h, space_before=20, space_after=8)
    run = h.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate 900
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    format_paragraph(h, space_before=14, space_after=6)
    run = h.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # Royal Blue
    return h

def add_heading_3(doc, text):
    h = doc.add_paragraph()
    format_paragraph(h, space_before=10, space_after=4)
    run = h.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69) # Slate 600
    return h

def add_body_p(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    format_paragraph(p, space_before=0, space_after=6)
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_bullet_point(doc, title, description=""):
    p = doc.add_paragraph(style='List Bullet')
    format_paragraph(p, space_before=2, space_after=4)
    run_t = p.add_run(title)
    run_t.font.name = 'Segoe UI'
    run_t.font.size = Pt(10.5)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    if description:
        run_d = p.add_run(f": {description}")
        run_d.font.name = 'Segoe UI'
        run_d.font.size = Pt(10.5)
        run_d.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_callout(doc, text, title="KEY ARCHITECTURAL HIGHLIGHT", border_hex="F97316", bg_hex="F8FAFC"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_hex}"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    p = cell.paragraphs[0]
    format_paragraph(p, space_before=6, space_after=6)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(10.5)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    run_text = p.add_run(text)
    run_text.font.name = 'Consolas' if border_hex in ["2563EB", "10B981"] else 'Segoe UI'
    run_text.font.size = Pt(9.5) if border_hex in ["2563EB", "10B981"] else Pt(10)
    run_text.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    empty_p = doc.add_paragraph()
    format_paragraph(empty_p, space_before=0, space_after=6)

def style_table_header(row, col_widths, bg_hex="0F172A"):
    for idx, cell in enumerate(row.cells):
        cell.width = Inches(col_widths[idx])
        set_cell_background(cell, bg_hex)
        set_cell_borders(cell, top="0F172A", bottom="0F172A", left="0F172A", right="0F172A")
        for p in cell.paragraphs:
            format_paragraph(p, space_before=4, space_after=4)
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def style_table_row(row, col_widths, is_even=False):
    bg_hex = "F8FAFC" if is_even else "FFFFFF"
    for idx, cell in enumerate(row.cells):
        cell.width = Inches(col_widths[idx])
        set_cell_background(cell, bg_hex)
        set_cell_borders(cell, top="E2E8F0", bottom="E2E8F0", left="E2E8F0", right="E2E8F0")
        for p in cell.paragraphs:
            format_paragraph(p, space_before=4, space_after=4)
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def build_single_complete_doc(filepath):
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    # Title Banner / Header
    title_p = doc.add_paragraph()
    format_paragraph(title_p, space_before=24, space_after=4)
    run_t = title_p.add_run("EduGuide 🎓")
    run_t.font.name = 'Segoe UI'
    run_t.font.size = Pt(32)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate 900
    
    sub_p = doc.add_paragraph()
    format_paragraph(sub_p, space_before=0, space_after=14)
    run_s = sub_p.add_run("Abstracted Student Learning Path Recommender & AI EdTech Platform\nComprehensive Master System Architecture & Technical Specifications")
    run_s.font.name = 'Segoe UI'
    run_s.font.size = Pt(14)
    run_s.font.italic = True
    run_s.font.color.rgb = RGBColor(0xF9, 0x73, 0x16) # Orange accent
    
    meta_p = doc.add_paragraph()
    format_paragraph(meta_p, space_before=0, space_after=20)
    run_m = meta_p.add_run("Document Version: 1.0.0 (Unified Master)  |  Stack: Next.js 16 + React 19 + PostgreSQL + OpenRouter AI")
    run_m.font.name = 'Segoe UI'
    run_m.font.size = Pt(9.5)
    run_m.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    # Horizontal Divider Bar
    divider = doc.add_table(rows=1, cols=1)
    divider.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = divider.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F97316")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    doc.add_paragraph()

    # Table of Contents Overview
    add_heading_1(doc, "Table of Contents Overview")
    toc_items = [
        ("1. Executive Summary & Core Vision", "Problem statement, solution concept, and skill node graph paradigm."),
        ("2. Requirement Fulfillment Matrix", "Core requirements mapped directly to architectural implementations."),
        ("3. Architecture & Technology Stack Specifications", "Frontend, Backend, Database, AI layers, and system topology."),
        ("4. Relational Database Schema & Data Dictionary", "PostgreSQL database tables, relationships, constraints, and SQL specs."),
        ("5. AI Recommendation Engine & Diagnostic Algorithms", "OpenRouter multi-model resiliency fallback chain and prompt engineering."),
        ("6. Frontend Component Architecture & Analytics Suite", "Next.js dashboard views, visual interactive components, and Recharts charts."),
        ("7. REST API Endpoint Specifications & Integration Contracts", "Endpoints, authentication headers, request/response JSON schemas."),
        ("8. Setup, Installation & Production Deployment Guide", "Prerequisites, environment configuration, quickstart commands, and deployment.")
    ]
    for section_title, section_desc in toc_items:
        add_bullet_point(doc, section_title, section_desc)

    doc.add_paragraph()

    # SECTION 1
    add_heading_1(doc, "1. Executive Summary & Core Vision")
    add_body_p(doc, "EduGuide is a state-of-the-art AI-powered EdTech ecosystem built to solve the modern challenge of complex, fragmented technical education. Traditional learning platforms present linear course lists that fail to adapt to a student's prior knowledge, target career goals, or weekly study schedule. EduGuide abstracts high-level career paths into granular, interconnected skill graph nodes.")
    add_body_p(doc, "By combining an AI Diagnostic Engine with interactive radar charts, real-time analytics, and adaptive path visualizers, EduGuide delivers an individualized learning trajectory that dynamically adjusts match confidence, milestone progress, and skill proficiency vectors.")

    add_callout(doc, 
        "EduGuide models learning domains as graph nodes (e.g., Spring Boot REST -> PostgreSQL JPA -> Next.js 16 -> Docker & Deployment), enabling personalized skill assessment, diagnostic gap analysis, and dynamic path recommendations.",
        title="CORE PARADIGM: SKILL GRAPH NODE ABSTRACTION"
    )

    # SECTION 2
    add_heading_1(doc, "2. Requirement Fulfillment Matrix")
    add_body_p(doc, "The table below outlines how each core requirement is solved by EduGuide's technical components:")

    matrix_table = doc.add_table(rows=5, cols=3)
    matrix_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [1.8, 2.2, 2.5]
    
    headers = ["Requirement Area", "Fulfillment Architecture", "Key Impact & Outcome"]
    for i, h in enumerate(headers):
        matrix_table.cell(0, i).paragraphs[0].text = h
    style_table_header(matrix_table.rows[0], col_widths)

    matrix_data = [
        ("Skill Abstraction", "Skill Graph Node Model & Skill Taxonomy Explorer", "Deconstructs complex software paths into bite-sized, sequential skill nodes with mastery level metrics."),
        ("AI Personalization", "OpenRouter AI Integration & Diagnostic Engine", "Evaluates diagnostic score, weekly commitment, and baseline answers to dynamically match ideal track (up to 98% fit)."),
        ("Progress Analytics", "Recharts Multi-Chart Dashboard System", "Visualizes daily activity velocity (AreaChart), skill radar (RadarChart), module balance (BarChart), and domain time (PieChart)."),
        ("Enterprise Architecture", "Next.js 16 App Router & Spring Boot REST API", "Delivers ultra-fast server rendering, client interactive charts, and production-ready REST synchronization.")
    ]

    for row_idx, data in enumerate(matrix_data, start=1):
        row = matrix_table.rows[row_idx]
        for col_idx, cell_value in enumerate(data):
            row.cells[col_idx].paragraphs[0].text = cell_value
        style_table_row(row, col_widths, is_even=(row_idx % 2 == 0))

    doc.add_paragraph()

    # SECTION 3
    add_heading_1(doc, "3. Architecture & Technology Stack Specifications")
    add_body_p(doc, "The EduGuide platform utilizes a modern hybrid architecture built on Next.js 16 (App Router) and React 19, backed by PostgreSQL and external AI LLM providers through OpenRouter.")

    add_heading_2(doc, "3.1 Technology Stack Specifications")
    add_bullet_point(doc, "Frontend Framework", "Next.js 16 (App Router with TypeScript & React 19)")
    add_bullet_point(doc, "Styling & UI Components", "Tailwind CSS v4, Lucide React Icon Library, Custom Design Tokens")
    add_bullet_point(doc, "Data Visualization", "Recharts 3.x (AreaChart, RadarChart, BarChart, PieChart)")
    add_bullet_point(doc, "Animation & Interactions", "Framer Motion 13.x micro-interactions and modal transitions")
    add_bullet_point(doc, "AI & LLM Services", "OpenRouter API multi-model fallback chain (Gemini 2.0 Flash, GPT-4o Mini, Llama 3.3 70B)")
    add_bullet_point(doc, "Database & Backend", "PostgreSQL database schema with Railway deployment ready REST endpoint integration (`https://eduguider.up.railway.app/api`)")

    add_heading_2(doc, "3.2 Layered System Topology")
    add_body_p(doc, "1. Client Layer: High-performance React 19 client components delivering interactive radar charts, path visualizers, modal generators, and diagnostic flow wizards.")
    add_body_p(doc, "2. API & Gateway Layer: Next.js API Routes (`/api/onboarding/recommend-skills`) handling server-side AI prompt synthesis, JSON schema validation, and fallback resilience.")
    add_body_p(doc, "3. Service & Repository Layer: `lib/api.ts` providing typed interfaces for JWT parsing, learning path fetching, user lesson progress updating, and transaction logging.")
    add_body_p(doc, "4. Database Layer: Relational PostgreSQL instance hosting core entities (`users`, `learning_paths`, `modules`, `lessons`, `user_skills`, `activity_logs`).")

    # SECTION 4
    add_heading_1(doc, "4. Relational Database Schema & Data Dictionary")
    add_body_p(doc, "The database schema is designed in PostgreSQL with strict data integrity, foreign key cascading, enum types, and check constraints to guarantee precision across all learning activities.")

    add_heading_2(doc, "4.1 PostgreSQL Enums Definition")
    add_bullet_point(doc, "user_role", "'STUDENT', 'ADMIN'")
    add_bullet_point(doc, "transaction_type", "'CREDIT', 'DEBIT'")
    add_bullet_point(doc, "difficulty_level", "'BEGINNER', 'INTERMEDIATE', 'ADVANCED'")
    add_bullet_point(doc, "progress_status", "'NOT_STARTED', 'IN_PROGRESS', 'COMPLETED'")

    add_heading_2(doc, "4.2 Summary of Relational Tables")
    
    db_table = doc.add_table(rows=9, cols=4)
    db_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    db_widths = [1.5, 1.2, 1.3, 2.5]
    
    db_headers = ["Table Name", "Primary Key", "Foreign Keys", "Core Responsibilities"]
    for i, h in enumerate(db_headers):
        db_table.cell(0, i).paragraphs[0].text = h
    style_table_header(db_table.rows[0], db_widths)

    db_rows_data = [
        ("users", "id (UUID)", "None", "Stores student profile, credentials hash, and system role."),
        ("transactions", "id (UUID)", "user_id -> users", "Tracks wallet transactions, course purchases, and credits."),
        ("learning_paths", "id (UUID)", "None", "Defines target career learning tracks, difficulty, and total hours."),
        ("modules", "id (UUID)", "None", "Contains core curriculum topics, primary video links, and durations."),
        ("lessons", "id (UUID)", "module_id -> modules", "Granular video lessons within modules, resource URLs, summaries."),
        ("path_modules", "id (UUID)", "path_id, module_id", "Junction table establishing ordered sequence of modules in paths."),
        ("user_learning_paths", "id (UUID)", "user_id, path_id", "Stores user active track, dynamic match score %, and completion %."),
        ("user_skills", "id (UUID)", "user_id, skill_id", "Tracks student skill mastery scores on a 1 to 10 radar vector scale.")
    ]

    for r_idx, r_data in enumerate(db_rows_data, start=1):
        row = db_table.rows[r_idx]
        for c_idx, val in enumerate(r_data):
            row.cells[c_idx].paragraphs[0].text = val
        style_table_row(row, db_widths, is_even=(r_idx % 2 == 0))

    doc.add_paragraph()

    add_heading_2(doc, "4.3 Data Dictionary Table Specifications")

    tables_spec = [
        ("users", [
            ("id", "UUID", "PRIMARY KEY", "gen_random_uuid()", "Unique student identifier"),
            ("name", "VARCHAR(255)", "NOT NULL", "-", "Full name of the user"),
            ("email", "VARCHAR(255)", "UNIQUE, NOT NULL", "-", "Login email address"),
            ("password_hash", "VARCHAR(255)", "NOT NULL", "-", "Bcrypt hashed password"),
            ("role", "user_role", "DEFAULT 'STUDENT'", "'STUDENT'", "User access permission role"),
            ("created_at", "TIMESTAMP W/ TZ", "DEFAULT CURRENT_TIMESTAMP", "NOW()", "Account creation timestamp")
        ]),
        ("learning_paths", [
            ("id", "UUID", "PRIMARY KEY", "gen_random_uuid()", "Unique path identifier"),
            ("title", "VARCHAR(255)", "NOT NULL", "-", "Career path title (e.g. Full-Stack Developer)"),
            ("description", "TEXT", "NULLABLE", "-", "Detailed summary of learning path"),
            ("level", "difficulty_level", "DEFAULT 'BEGINNER'", "'BEGINNER'", "Overall track difficulty tier"),
            ("estimated_hours", "INTEGER", "NOT NULL", "-", "Total estimated completion hours")
        ]),
        ("modules", [
            ("id", "UUID", "PRIMARY KEY", "gen_random_uuid()", "Unique module identifier"),
            ("title", "VARCHAR(255)", "NOT NULL", "-", "Module subject title"),
            ("topic", "VARCHAR(100)", "NOT NULL", "-", "Domain category (e.g. Backend, Frontend)"),
            ("duration_minutes", "INTEGER", "NOT NULL", "-", "Estimated module duration in minutes"),
            ("video_url", "TEXT", "NULLABLE", "-", "Primary video overview URL")
        ]),
        ("lessons", [
            ("id", "UUID", "PRIMARY KEY", "gen_random_uuid()", "Unique lesson identifier"),
            ("module_id", "UUID", "FK -> modules(id)", "-", "Module parent reference"),
            ("title", "VARCHAR(255)", "NOT NULL", "-", "Lesson title"),
            ("sequence_order", "INTEGER", "NOT NULL", "-", "Order of lesson inside module"),
            ("resources_url", "TEXT", "NULLABLE", "-", "Downloadable resource or code link")
        ]),
        ("user_learning_paths", [
            ("id", "UUID", "PRIMARY KEY", "gen_random_uuid()", "Unique junction record ID"),
            ("user_id", "UUID", "FK -> users(id)", "-", "Cascades on delete"),
            ("path_id", "UUID", "FK -> learning_paths(id)", "-", "Cascades on delete"),
            ("is_active", "BOOLEAN", "DEFAULT false", "false", "Current active enrolled track"),
            ("match_score", "INTEGER", "CHECK (0..100)", "-", "AI calculated match score percentage"),
            ("progress_percentage", "INTEGER", "CHECK (0..100)", "0", "Overall track completion percentage")
        ]),
        ("user_skills", [
            ("id", "UUID", "PRIMARY KEY", "gen_random_uuid()", "Unique skill vector record"),
            ("user_id", "UUID", "FK -> users(id)", "-", "Student user ID"),
            ("skill_id", "UUID", "FK -> skills(id)", "-", "Skill reference ID"),
            ("mastery_level", "INTEGER", "CHECK (1..10)", "-", "Radar score proficiency level (1 to 10)")
        ])
    ]

    for t_name, fields in tables_spec:
        add_heading_3(doc, f"Table: {t_name}")
        t = doc.add_table(rows=len(fields)+1, cols=5)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        widths = [1.2, 1.1, 1.4, 1.2, 1.6]
        
        headers = ["Column", "Data Type", "Constraints", "Default", "Description"]
        for idx, h in enumerate(headers):
            t.cell(0, idx).paragraphs[0].text = h
        style_table_header(t.rows[0], widths, bg_hex="1E293B")
        
        for r_idx, f in enumerate(fields, start=1):
            row = t.rows[r_idx]
            for c_idx, val in enumerate(f):
                row.cells[c_idx].paragraphs[0].text = val
            style_table_row(row, widths, is_even=(r_idx % 2 == 0))
            
        doc.add_paragraph()

    # SECTION 5
    add_heading_1(doc, "5. AI Recommendation Engine & Diagnostic Algorithms")
    add_body_p(doc, "The EduGuide AI Engine bridges the gap between student diagnostic assessment and personalized curriculum planning.")

    add_heading_2(doc, "5.1 Multi-Model AI Resiliency Pattern")
    add_body_p(doc, "The AI infrastructure (`lib/openrouter.ts`) features an automatic fallback strategy across high-performing LLM models to maintain 99.99% availability even under API rate limits or budget constraints:")
    add_bullet_point(doc, "Primary Model", "Google Gemini 2.0 Flash Lite (`google/gemini-2.0-flash-lite-001`)")
    add_bullet_point(doc, "Secondary Fallbacks", "OpenAI GPT-4o Mini, Meta Llama 3.1 8B Instruct, Qwen 2.5 72B, DeepSeek Chat, Llama 3.3 70B")
    add_bullet_point(doc, "Dynamic Token Budget Retry", "Automatically parses HTTP 402 budget warnings and retries requests at exact affordable token limits down to 300 tokens.")

    add_heading_2(doc, "5.2 Recommendation Data Flow")
    add_body_p(doc, "1. Student completes onboarding diagnostic quiz selecting technical proficiency and weekly hours commitment.")
    add_body_p(doc, "2. POST request dispatched to `/api/onboarding/recommend-skills` containing diagnostic scores and available track catalog.")
    add_body_p(doc, "3. AI model synthesizes diagnostic answers, calculates fit percentages, and generates custom student strength highlights.")
    add_body_p(doc, "4. Response updates user dashboard with active track recommendation, unlocking personalized roadmap visualizers.")

    # SECTION 6
    add_heading_1(doc, "6. Frontend Component Architecture & Analytics Suite")
    add_body_p(doc, "EduGuide offers an enterprise-class UI component framework designed around clarity, rapid feedback, and engaging visual aesthetics.")

    add_heading_2(doc, "6.1 Core Dashboard Views & Visualizations")
    add_bullet_point(doc, "Overview View (`OverviewView.tsx`)", "Central hub displaying active path progress, overall score cards, active module visualizer, and quick action modals.")
    add_bullet_point(doc, "Analytics Dashboard (`AnalyticsCharts.tsx`)", "Contains 4 dynamic Recharts components: Weekly Activity AreaChart, Skill Competency RadarChart, Module BarChart, and Domain PieChart.")
    add_bullet_point(doc, "Roadmap Visualizer (`PathVisualizer.tsx`)", "Interactive node step visualizer with milestone toggles, status pills, estimated time indicators, and completion checkmarks.")
    add_bullet_point(doc, "Path Generator Modal (`GeneratePathModal.tsx`)", "Framer-Motion animated modal allowing students to request AI custom learning path generation.")
    add_bullet_point(doc, "Skill Explorer (`ExplorerView.tsx`)", "Taxonomy explorer categorized into Backend Architecture, Database Systems, Frontend Frameworks, and DevOps & Security.")

    # SECTION 7
    add_heading_1(doc, "7. REST API Endpoint Specifications & Integration Contracts")
    add_body_p(doc, "The backend integration endpoints are standardized to communicate via REST JSON formats.")

    api_table = doc.add_table(rows=6, cols=4)
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_widths = [1.2, 1.8, 1.5, 2.0]
    
    api_headers = ["HTTP Method", "Endpoint Path", "Request / Payload", "Functionality Description"]
    for i, h in enumerate(api_headers):
        api_table.cell(0, i).paragraphs[0].text = h
    style_table_header(api_table.rows[0], api_widths)

    api_rows_data = [
        ("POST", "/api/onboarding/recommend-skills", "Diagnostic answers, weekly hours, skill level", "Invokes OpenRouter AI to return recommended track ID, match %, and strengths."),
        ("GET", "/api/learning-paths", "Authorization Bearer JWT", "Fetches all available learning path tracks with metadata and module counts."),
        ("GET", "/api/user/active-path", "Authorization Bearer JWT", "Retrieves active user learning path, match score, and module progress percentage."),
        ("PUT", "/api/user/lessons/{id}/progress", "Status: COMPLETED / IN_PROGRESS", "Updates specific lesson progress and updates student module completion stats."),
        ("GET", "/api/dashboard/analytics", "Authorization Bearer JWT", "Aggregates student weekly activity logs, skill mastery scores, and time distribution.")
    ]

    for r_idx, r_data in enumerate(api_rows_data, start=1):
        row = api_table.rows[r_idx]
        for c_idx, val in enumerate(r_data):
            row.cells[c_idx].paragraphs[0].text = val
        style_table_row(row, api_widths, is_even=(r_idx % 2 == 0))

    doc.add_paragraph()

    add_heading_2(doc, "7.1 AI Skill Recommendation API Contract")
    add_heading_3(doc, "POST /api/onboarding/recommend-skills")
    
    req_json = (
        '{\n'
        '  "skillLevel": "INTERMEDIATE",\n'
        '  "weeklyHours": 10,\n'
        '  "diagnosticScore": 75,\n'
        '  "diagnosticAnswers": {\n'
        '    "q1": 4,\n'
        '    "q2": 5,\n'
        '    "q3": 3\n'
        '  }\n'
        '}'
    )
    add_callout(doc, req_json, title="SAMPLE REQUEST PAYLOAD", border_hex="2563EB", bg_hex="F1F5F9")

    res_json = (
        '{\n'
        '  "success": true,\n'
        '  "recommendation": {\n'
        '    "recommendedTrackId": "track-fullstack-spring-next",\n'
        '    "recommendedTrackTitle": "Full-Stack Spring Boot & Next.js Track",\n'
        '    "matchScorePercent": 96,\n'
        '    "recommendedLevelTier": "INTERMEDIATE",\n'
        '    "strengths": [\n'
        '      "Strong backend database normalization understanding",\n'
        '      "Solid familiarity with React component lifecycles"\n'
        '    ]\n'
        '  }\n'
        '}'
    )
    add_callout(doc, res_json, title="SAMPLE SUCCESS RESPONSE PAYLOAD", border_hex="10B981", bg_hex="F0FDF4")

    # SECTION 8
    add_heading_1(doc, "8. Installation, Configuration & Production Deployment Guide")
    add_body_p(doc, "Follow the steps below to initialize and deploy the EduGuide project in local development or production environments.")

    add_heading_2(doc, "8.1 Prerequisites")
    add_bullet_point(doc, "Node.js Environment", "Node.js v20.x or higher, npm v10.x+")
    add_bullet_point(doc, "Database Engine", "PostgreSQL 15+ instance with `gen_random_uuid()` extension support.")
    add_bullet_point(doc, "API Credentials", "OpenRouter API Key (`OPEN_ROUTER_API_KEY`) for AI recommendations.")

    add_heading_2(doc, "8.2 Quickstart Execution Commands")
    
    cmd_table = doc.add_table(rows=4, cols=2)
    cmd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cmd_widths = [2.2, 4.3]
    
    cmd_table.cell(0, 0).paragraphs[0].text = "Step / Stage"
    cmd_table.cell(0, 1).paragraphs[0].text = "Command Line Execution"
    style_table_header(cmd_table.rows[0], cmd_widths)

    cmd_data = [
        ("1. Install Dependencies", "npm install"),
        ("2. Environment Configuration", "cp .env.example .env.local\n# Configure OPEN_ROUTER_API_KEY and NEXT_PUBLIC_BACKEND_URL"),
        ("3. Launch Dev Server", "npm run dev  # Opens http://localhost:3000")
    ]

    for r_idx, r_data in enumerate(cmd_data, start=1):
        row = cmd_table.rows[r_idx]
        for c_idx, val in enumerate(r_data):
            row.cells[c_idx].paragraphs[0].text = val
        style_table_row(row, cmd_widths, is_even=(r_idx % 2 == 0))

    doc.add_paragraph()
    
    add_callout(doc,
        "EduGuide is configured for zero-friction production deployment on Railway or Vercel. Database schema migrations can be executed directly using schema.sql.",
        title="PRODUCTION DEPLOYMENT READY"
    )

    doc.save(filepath)
    print(f"Successfully generated master doc: {filepath}")

if __name__ == '__main__':
    base_dir = os.path.dirname(os.path.abspath(__file__))
    master_path = os.path.join(base_dir, "EduGuide_Complete_Project_Documentation.docx")
    build_single_complete_doc(master_path)
